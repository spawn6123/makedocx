from docx import Document
from docx.shared import Cm
from PIL import Image
import os


def createdoc(foldername, currentpath, filenameBase, images, savepath):
    """
    製作Docx另存資料匣名稱
    """
    demofile = os.path.join(currentpath, 'demo.docx')
    document = Document(demofile)
    for i in images:
        ig = Image.open(i)
        w, h = ig.size
        if w > h:
            document.add_picture(i, width=Cm(14.0))
        else:
            document.add_picture(i, height=Cm(20.0))
    base = os.path.basename(filenameBase)
    foldername = os.path.basename(foldername)
    outfile = os.path.join(savepath, foldername + '-' + base + '.docx')
    # print(outfile)
    document.save(outfile)
